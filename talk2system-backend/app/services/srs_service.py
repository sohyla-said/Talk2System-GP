import os
import re
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.db.session import get_db
from app.services.session_service import SessionService
from app.db.session import SessionLocal
from app.models.background_task import BackgroundTask
from app.models.project import Project
from app.models.session_membership import SessionMembership
from app.services import notification_service
import logging
from app.services.audit_service import log_action
logger = logging.getLogger(__name__)
# ==========================
# CONFIG
# ==========================
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "qwen2.5:7b"
STORAGE_PATH = "storage/srs"

# ==========================
# SUPPORTED FORMAT VERSIONS
# ==========================
SUPPORTED_FORMATS = ["ieee_830", "iso_iec_29148", "modern_agile"]

# ==========================
# HELPERS
# ==========================
def format_requirements_for_prompt(requirements_json: dict) -> str:
    actors = requirements_json.get("actors", [])
    features = requirements_json.get("features", [])
    frs = requirements_json.get("functional_requirements", [])
    nfrs = requirements_json.get("nonfunctional_requirements", []) or \
           requirements_json.get("non_functional_requirements", [])

    lines = []

    if actors:
        lines.append("ACTORS:")
        for a in actors:
            lines.append(f"  - {a}")

    if features:
        lines.append("\nFEATURES:")
        for f in features:
            lines.append(f"  - {f}")

    if frs:
        lines.append("\nFUNCTIONAL REQUIREMENTS:")
        for i, fr in enumerate(frs, 1):
            actor = fr.get("actor", "")
            text = fr.get("text", "")
            feature = fr.get("feature", "")
            lines.append(f"  FR-{i}: [{actor}] {text} (Feature: {feature})")

    if nfrs:
        lines.append("\nNON-FUNCTIONAL REQUIREMENTS:")
        for i, nfr in enumerate(nfrs, 1):
            category = nfr.get("category", "General")
            text = nfr.get("text", "")
            lines.append(f"  NFR-{i}: [{category}] {text}")

    return "\n".join(lines)

# ==========================
# REQUIREMENTS CHUNKING
# ==========================
# Threshold: if total FRs + NFRs exceed this, split into chunks
FR_CHUNK_THRESHOLD = 10   # anything above 10 FRs+NFRs gets chunked
FR_CHUNK_SIZE      = 8    # 8 FRs per chunk — small enough for CPU
NFR_CHUNK_SIZE     = 5    # 5 NFRs per chunk

def chunk_requirements(requirements_json: dict) -> list[dict]:
    """
    Splits a large requirements_json into smaller chunks, each under
    FR_CHUNK_SIZE functional requirements and NFR_CHUNK_SIZE non-functional.
    Returns a list of requirements_json dicts — each is a valid slice.
    If requirements are small enough, returns a single-element list (no chunking).
    """
    frs  = requirements_json.get("functional_requirements", [])
    nfrs = requirements_json.get("nonfunctional_requirements", []) or \
           requirements_json.get("non_functional_requirements", [])
    actors   = requirements_json.get("actors", [])
    features = requirements_json.get("features", [])

    total = len(frs) + len(nfrs)

    # ── No chunking needed ─────────────────────────────────────────────────
    if total <= FR_CHUNK_THRESHOLD:
        logger.info(f"Requirements within threshold ({total} items) — no chunking needed")
        return [requirements_json]

    logger.info(
        f"Requirements exceed threshold ({len(frs)} FRs, {len(nfrs)} NFRs) "
        f"— splitting into chunks of {FR_CHUNK_SIZE} FRs / {NFR_CHUNK_SIZE} NFRs"
    )

    chunks = []
    fr_batches  = [frs[i:i + FR_CHUNK_SIZE]  for i in range(0, max(len(frs), 1),  FR_CHUNK_SIZE)]
    nfr_batches = [nfrs[i:i + NFR_CHUNK_SIZE] for i in range(0, max(len(nfrs), 1), NFR_CHUNK_SIZE)]

    # Pad the shorter list so zip covers everything
    max_batches = max(len(fr_batches), len(nfr_batches))
    while len(fr_batches)  < max_batches: fr_batches.append([])
    while len(nfr_batches) < max_batches: nfr_batches.append([])

    for idx, (fr_batch, nfr_batch) in enumerate(zip(fr_batches, nfr_batches)):
        chunk = {
            "actors":   actors,    # keep actors + features in every chunk for context
            "features": features,
            "functional_requirements":     fr_batch,
            "nonfunctional_requirements":  nfr_batch,
            "_chunk_index": idx + 1,       # internal metadata for logging
            "_total_chunks": max_batches,
        }
        chunks.append(chunk)
        logger.info(
            f"Chunk {idx+1}/{max_batches}: "
            f"{len(fr_batch)} FRs, {len(nfr_batch)} NFRs"
        )

    return chunks


def _call_ollama_for_chunk(
    prompt: str,
    chunk_index: int,
    total_chunks: int,
    timeout: int = 600,
) -> str:
    """
    Calls Ollama with a single SRS chunk prompt.
    Raises on connection error or non-200 status.
    """
    logger.info(f"Calling Ollama for SRS chunk {chunk_index}/{total_chunks}...")
    try:
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.3,
                    "num_predict": 800,  # per chunk — enough for partial sections
                    "num_ctx":    2048,   # reduced context = faster CPU inference
                }
            },
            timeout=timeout,
        )
    except requests.exceptions.ConnectionError:
        raise Exception("Ollama is not running. Start it with: ollama serve")
    except requests.exceptions.Timeout:
        raise Exception(
            f"Ollama timed out on chunk {chunk_index}/{total_chunks}. "
            "Try reducing requirements or restarting Ollama."
        )

    if response.status_code != 200:
        raise Exception(
            f"Ollama error on chunk {chunk_index}/{total_chunks}: "
            f"{response.status_code} — {response.text[:200]}"
        )

    text = response.json().get("response", "").strip()
    logger.info(
        f"Chunk {chunk_index}/{total_chunks} done — "
        f"{len(text.split())} words generated"
    )
    return text

def build_chunk_prompt(
    chunk: dict,
    project_name: str,
    chunk_index: int,
    total_chunks: int,
    format_version: str,
) -> str:
    formatted = format_requirements_for_prompt(chunk)

    if chunk_index == 1:
        return f"""You are a senior software engineer writing a formal SRS document for "{project_name}".
This is part {chunk_index} of {total_chunks}. Write the document introduction and the requirements below.

Requirements Data:
{formatted}

Output ONLY using this structure:

# 1. Introduction
## 1.1 Purpose
## 1.2 Scope

# 2. Overall Description
## 2.1 Product Perspective
## 2.2 User Classes

# 3. Functional Requirements
For each FR use exactly this format on two lines:
**FR-N** | Actor: [actor] | Feature: [feature]
The system shall [requirement text].

# 4. Non-Functional Requirements
For each NFR use exactly this format:
**NFR-N** | Category: [category]
[requirement text].

STRICT RULES:
- Output ONLY the document, nothing before or after it
- Start directly with # 1. Introduction
- Use "shall" for all functional requirements
- Do NOT include traceability matrices or acceptance criteria sections
- Do NOT skip any requirement from the data above
"""

    # ── Chunks 2+: output ONLY FR and NFR lines, nothing else ─────────────
    frs  = chunk.get("functional_requirements", [])
    nfrs = chunk.get("nonfunctional_requirements", []) or \
           chunk.get("non_functional_requirements", [])

    fr_start  = (chunk_index - 1) * FR_CHUNK_SIZE + 1
    nfr_start = (chunk_index - 1) * NFR_CHUNK_SIZE + 1

    # Build the FR/NFR entries directly — no LLM needed for continuation chunks
    # This guarantees ALL requirements appear and avoids Ollama hallucinating structure
    lines = []
    for i, fr in enumerate(frs, fr_start):
        actor   = fr.get("actor", "system")
        text    = fr.get("text", "")
        feature = fr.get("feature", "")
        lines.append(f"**FR-{i}** | Actor: {actor} | Feature: {feature}")
        lines.append(f"The system shall {text.lower().lstrip('the system must').lstrip('the system shall').lstrip('users must').strip()}.")
        lines.append("")

    for i, nfr in enumerate(nfrs, nfr_start):
        category = nfr.get("category", "General")
        text     = nfr.get("text", "")
        lines.append(f"**NFR-{i}** | Category: {category}")
        lines.append(f"{text}")
        lines.append("")

    return "\n".join(lines)

# ==========================
# BUILD PROMPT — IEEE 830
# (original, preserved exactly)
# ==========================
def build_srs_prompt_ieee_830(requirements_json: dict, project_name: str) -> str:
    formatted = format_requirements_for_prompt(requirements_json)

    return f"""You are a senior software engineer writing a formal IEEE 830 Software Requirements Specification (SRS).

Generate a complete SRS document for the system named "{project_name}".

Requirements Data:
{formatted}

Output ONLY the document using this EXACT structure with these EXACT section headings:

# 1. Introduction
## 1.1 Purpose
## 1.2 Scope
## 1.3 Definitions and Acronyms

# 2. Overall Description
## 2.1 Product Perspective
## 2.2 User Classes and Characteristics
## 2.3 Assumptions and Dependencies

# 3. Functional Requirements
List every functional requirement using this format for each one:
**FR-N** | Actor: [actor] | Feature: [feature]
The system shall [requirement text].

# 4. Non-Functional Requirements
## 4.1 Performance
## 4.2 Security
## 4.3 Usability
## 4.4 Reliability

STRICT RULES:
- Output ONLY the document, nothing before or after it
- Start directly with # 1. Introduction
- Do NOT include a section for External Interfaces
- Do NOT include a Traceability Matrix
- Use "shall" for all functional requirements
- Keep NFRs measurable with numbers where possible
"""


# ==========================
# BUILD PROMPT — ISO/IEC/IEEE 29148
# ==========================
def build_srs_prompt_iso_29148(requirements_json: dict, project_name: str) -> str:
    formatted = format_requirements_for_prompt(requirements_json)

    return f"""You are a senior systems engineer writing a formal Software Requirements Specification (SRS) compliant with ISO/IEC/IEEE 29148:2018.

Generate a complete SRS document for the system named "{project_name}".

Requirements Data:
{formatted}

Output ONLY the document using this EXACT structure with these EXACT section headings:

# 1. Introduction
## 1.1 Purpose
## 1.2 Scope
## 1.3 Definitions, Acronyms, and Abbreviations
## 1.4 References
## 1.5 Overview

# 2. Overall Description
## 2.1 Product Perspective
## 2.2 Product Functions
## 2.3 User Characteristics
## 2.4 Constraints
## 2.5 Assumptions and Dependencies

# 3. Specific Requirements
## 3.1 Functional Requirements
List every functional requirement using this format:
**FR-N** | Actor: [actor] | Feature: [feature]
The system shall [requirement text].

## 3.2 Non-Functional Requirements
### 3.2.1 Performance Requirements
### 3.2.2 Security Requirements
### 3.2.3 Usability Requirements
### 3.2.4 Reliability Requirements
### 3.2.5 Maintainability Requirements

# 4. Verification
## 4.1 Verification Methods
## 4.2 Acceptance Criteria
For each major functional requirement group, state the verification method (Test, Inspection, Demonstration, Analysis).

# 5. Traceability
## 5.1 Requirements Traceability Matrix
List each FR with its source actor, feature, and verification method in a table format:
| Req ID | Description Summary | Actor | Feature | Verification |

STRICT RULES:
- Output ONLY the document, nothing before or after it
- Start directly with # 1. Introduction
- Use "shall" for all mandatory functional requirements
- Use "should" for recommended non-functional requirements
- Keep NFRs measurable with specific metrics where possible
- Include at least one verification method per functional requirement group
"""


# ==========================
# BUILD PROMPT — Modern Agile SRS
# ==========================
def build_srs_prompt_modern_agile(requirements_json: dict, project_name: str) -> str:
    formatted = format_requirements_for_prompt(requirements_json)

    return f"""You are a senior product engineer writing a Modern Agile Software Requirements Specification (SRS).
This format prioritizes clarity, traceability, and user-centric language over rigid formality.

Generate a complete SRS document for the system named "{project_name}".

Requirements Data:
{formatted}

Output ONLY the document using this EXACT structure with these EXACT section headings:

# 1. Introduction
## 1.1 Purpose and Vision
## 1.2 Scope
## 1.3 Stakeholders and Users
## 1.4 Glossary

# 2. System Context
## 2.1 Product Overview
## 2.2 Key User Personas
## 2.3 Constraints and Assumptions

# 3. Functional Requirements
For each functional requirement, use this format:
**FR-N** | Actor: [actor] | Feature: [feature]
As a [actor], I need to [requirement text] so that [business value].

# 4. Non-Functional Requirements
## 4.1 Performance
## 4.2 Security and Privacy
## 4.3 Usability and Accessibility
## 4.4 Reliability and Availability
## 4.5 AI and Automation Considerations

# 5. Acceptance Criteria
For each major feature group, list 2-3 measurable acceptance criteria using Given/When/Then format:
**Feature: [feature name]**
- Given [context], When [action], Then [expected outcome].

# 6. Traceability
## 6.1 Requirements Traceability
| Req ID | User Story Summary | Feature | Priority | Acceptance Criteria Ref |

STRICT RULES:
- Output ONLY the document, nothing before or after it
- Start directly with # 1. Introduction
- Use "must" for mandatory requirements, "should" for recommended ones
- Write functional requirements as user stories (As a / I need to / so that)
- Keep acceptance criteria concrete and testable
- Include AI/automation notes where relevant to modern systems
"""


# ==========================
# ROUTER — picks the right prompt builder
# ==========================
def build_srs_prompt(requirements_json: dict, project_name: str, format_version: str = "ieee_830") -> str:
    if format_version == "iso_iec_29148":
        return build_srs_prompt_iso_29148(requirements_json, project_name)
    elif format_version == "modern_agile":
        return build_srs_prompt_modern_agile(requirements_json, project_name)
    else:
        # default: ieee_830 (original behavior preserved)
        return build_srs_prompt_ieee_830(requirements_json, project_name)


# ==========================
# CALL OLLAMA
# ==========================
def generate_srs_text_with_ollama(
    requirements_json: dict,
    project_name: str = "Software System",
    format_version: str = "ieee_830"
) -> str:
    """
    Generates SRS text via Ollama.
    - If requirements are small (≤ FR_CHUNK_THRESHOLD items): single prompt, single call.
    - If requirements are large: splits into chunks, calls Ollama per chunk,
      then merges all outputs into one coherent document.
    """
    chunks = chunk_requirements(requirements_json)

    if len(chunks) == 1:
        # ── Fast path: single prompt (original behavior) ───────────────────
        logger.info("Single-chunk SRS generation (no splitting needed)")
        prompt = build_srs_prompt(requirements_json, project_name, format_version)
        return _call_ollama_for_chunk(prompt, chunk_index=1, total_chunks=1)

    # ── Chunked path ───────────────────────────────────────────────────────
    logger.info(f"Chunked SRS generation: {len(chunks)} chunks")
    parts = []

    for chunk in chunks:
        idx   = chunk["_chunk_index"]
        total = chunk["_total_chunks"]

        prompt = build_chunk_prompt(
            chunk=chunk,
            project_name=project_name,
            chunk_index=idx,
            total_chunks=total,
            format_version=format_version,
        )

        if idx == 1:
            # Only chunk 1 goes to Ollama — it generates the doc skeleton + first batch
            part = _call_ollama_for_chunk(
                prompt=prompt,
                chunk_index=idx,
                total_chunks=total,
                timeout=600,
            )
            logger.info(f"SRS chunk {idx}/{total} from Ollama ({len(part.split())} words)")
        else:
            # Chunks 2+ are pre-built directly from requirements — no Ollama call
            # This guarantees ALL requirements appear, no timeout risk
            part = prompt   # prompt IS the content for continuation chunks
            logger.info(
                f"SRS chunk {idx}/{total} built directly "
                f"({len(part.split())} words, no Ollama call)"
            )

        parts.append(part)

    # ── Merge all parts ────────────────────────────────────────────────────
    merged = _merge_srs_chunks(parts)
    logger.info(
        f"SRS generation complete — {len(chunks)} chunks merged, "
        f"{len(merged.split())} total words"
    )
    return merged

def _merge_srs_chunks(parts: list[str]) -> str:
    """
    Merges multiple SRS chunk outputs into one document.
    Strategy: inject continuation FR/NFR entries into section 3/4 of part 1,
    not after section 4 — so all requirements appear in the right place.
    """
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0]

    # Collect all continuation content from chunks 2+
    continuation_frs  = []
    continuation_nfrs = []

    for part in parts[1:]:
        cleaned = part.strip()
        # Remove any accidental section 1/2 headers Ollama adds
        cleaned = re.sub(r"^(#\s*(1\.|2\.)[^\n]*\n)+", "", cleaned, flags=re.MULTILINE).strip()

        for line in cleaned.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            # Detect FR lines: **FR-N** | ... or FR-N | ...
            if re.match(r"^\*?\*?FR-\d+\*?\*?", stripped):
                continuation_frs.append(stripped)
            # Detect NFR lines: NFR-N | ... or **NFR-N** | ...
            elif re.match(r"^\*?\*?NFR-\d+\*?\*?", stripped):
                continuation_nfrs.append(stripped)
            # Detect "The system shall..." lines that follow an FR header
            elif stripped.lower().startswith("the system shall") or \
                 stripped.lower().startswith("the system must") or \
                 stripped.lower().startswith("as a "):
                # Belongs to the last FR entry
                if continuation_frs:
                    continuation_frs[-1] = continuation_frs[-1] + "\n" + stripped
                else:
                    continuation_frs.append(stripped)

    if not continuation_frs and not continuation_nfrs:
        # Nothing useful extracted — fall back to simple append
        merged = parts[0].rstrip()
        for i, part in enumerate(parts[1:], start=2):
            cleaned = part.strip()
            cleaned = re.sub(r"^(#\s*(1\.|2\.)[^\n]*\n)+", "", cleaned, flags=re.MULTILINE).strip()
            if cleaned:
                merged += f"\n\n{cleaned}"
        return merged

    base = parts[0]

    # ── Inject continuation FRs into section 3 (before section 4) ────────
    # Find the boundary between section 3 and section 4 in chunk 1
    # Section 4 starts with "# 4." or "## 4." or "# Non-Functional"
    section4_pattern = re.compile(
        r"(\n#+\s*(4\.|Non-Functional Requirements))",
        re.IGNORECASE
    )
    match = section4_pattern.search(base)

    if match and continuation_frs:
        insert_pos = match.start()
        fr_block = "\n" + "\n".join(continuation_frs) + "\n"
        base = base[:insert_pos] + fr_block + base[insert_pos:]
        logger.info(f"Injected {len(continuation_frs)} continuation FR entries into section 3")

    # ── Inject continuation NFRs at the end of section 4 ─────────────────
    if continuation_nfrs:
        nfr_block = "\n" + "\n".join(continuation_nfrs)
        base = base.rstrip() + nfr_block
        logger.info(f"Injected {len(continuation_nfrs)} continuation NFR entries into section 4")

    logger.info(f"Merge complete — {len(continuation_frs)} FRs + {len(continuation_nfrs)} NFRs injected")
    return base

# ==========================
# FORMAT DISPLAY NAMES (used in docx cover subtitle)
# ==========================
FORMAT_DISPLAY_NAMES = {
    "ieee_830": "IEEE Std 830 — Formal Requirements Document",
    "iso_iec_29148": "ISO/IEC/IEEE 29148:2018 — Systems and Software Requirements",
    "modern_agile": "Modern Agile SRS — User-Centric Requirements Specification",
}


# ==========================
# SAVE AS DOCX
# ==========================
def save_srs_as_docx(
    srs_text: str,
    project_id: int,
    session_id: int = None,
    format_version: str = "ieee_830"
) -> str:
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    os.makedirs(STORAGE_PATH, exist_ok=True)

    if session_id:
        filename = f"srs_P{project_id}_S{session_id}_{int(datetime.utcnow().timestamp())}.docx"
    else:
        filename = f"srs_P{project_id}_{int(datetime.utcnow().timestamp())}.docx"

    path = os.path.join(STORAGE_PATH, filename)
    doc = Document()

    # ---- Page margins (IEEE-style) ----
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ---- Cover block ----
    cover_title = doc.add_heading("Software Requirements Specification", level=0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.runs[0]
    run.font.color.rgb = RGBColor(0x1E, 0x10, 0x5F)
    run.font.size = Pt(24)

    # Use the format-specific subtitle
    subtitle_text = FORMAT_DISPLAY_NAMES.get(format_version, "IEEE Std 830 — Formal Requirements Document")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(subtitle_text).font.size = Pt(11)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')} UTC")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ---- Horizontal rule after cover ----
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E105F')
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()

    # ---- Parse markdown into styled Word content ----
    for line in srs_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(p, stripped[2:])
        elif re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_bold(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif stripped.startswith("|"):
            # simple table row — render as plain paragraph in docx
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            p = doc.add_paragraph()
            p.add_run(" | ".join(cells))
        elif stripped.startswith("---"):
            # horizontal rule → paragraph border
            hr_p = doc.add_paragraph()
            hr_pPr = hr_p._p.get_or_add_pPr()
            hr_pBdr = OxmlElement('w:pBdr')
            hr_bottom = OxmlElement('w:bottom')
            hr_bottom.set(qn('w:val'), 'single')
            hr_bottom.set(qn('w:sz'), '4')
            hr_bottom.set(qn('w:space'), '1')
            hr_bottom.set(qn('w:color'), 'CCCCCC')
            hr_pBdr.append(hr_bottom)
            hr_pPr.append(hr_pBdr)
        else:
            p = doc.add_paragraph()
            _add_inline_bold(p, stripped)

    doc.save(path)
    return path


# ==========================
# INLINE BOLD HELPER
# ==========================
def _add_inline_bold(paragraph, text: str):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ==========================
# MAIN PIPELINE
# ==========================
def generate_srs_pipeline(
    requirements_json: dict,
    project_id: int,
    project_name: str = "Software System",
    session_id: int = None,
    format_version: str = "ieee_830"   # ← NEW PARAM, defaults to original behavior
) -> dict:
    try:
        srs_text = generate_srs_text_with_ollama(requirements_json, project_name, format_version)
        file_path = save_srs_as_docx(srs_text, project_id, session_id, format_version)
        if session_id:
            db = next(get_db())
            try:
                SessionService.update_session_status(db, session_id, "pending approval")
            finally:
                db.close()
        return {
            "srs_text": srs_text,
            "file_path": file_path
        }
    except Exception as e:
        raise Exception(f"SRS generation failed: {str(e)}")
    
def _sanitize_srs_error_message(raw: str) -> str:
    """Convert raw exception text into a short user-friendly message."""
    low = raw.lower()
    if "timeout" in low or "timed out" in low:
        return "SRS generation timed out. Please ensure Ollama is running and retry."
    if "ollama" in low or "connection" in low:
        return "Could not connect to Ollama. Please ensure it is running and retry."
    if "quota" in low or "429" in raw:
        return "AI generation quota exceeded. Please retry in a few minutes."
    return "SRS generation failed. Please retry."

def run_async_srs_task(
    task_id: int,
    project_id: int,
    session_id: int | None,
    format_version: str,
    source: str,          # "session" or "project"
    user_id: int,
):
    db = SessionLocal()
    try:
        task = db.query(BackgroundTask).filter(BackgroundTask.id == task_id).first()
        task.status = "in-progress"
        db.commit()

        # ── 1. Fetch requirements ──────────────────────────────────────────
        from app.services.requirement_service import RequirementService
        if source == "session":
            req = RequirementService.get_latest_session_requirement(db, project_id, session_id)
        else:
            req = RequirementService.get_latest_project_requirement(db, project_id)
        requirements_json = req["data"]

        # ── 2. Project name ────────────────────────────────────────────────
        project = db.query(Project).filter(Project.id == project_id).first()
        project_name = project.name if project else "Software System"

        # ── 3. Run the pipeline ────────────────────────────────────────────
        result = generate_srs_pipeline(
            requirements_json=requirements_json,
            project_id=project_id,
            project_name=project_name,
            session_id=session_id if source == "session" else None,
            format_version=format_version,
        )

        # ── 4. Save artifact ───────────────────────────────────────────────
        from app.services.artifact_service import ArtifactService
        artifact = ArtifactService.save_artifact(
            db=db,
            project_id=project_id,
            session_id=session_id if source == "session" else None,
            artifact_type_name="SRS_DOCUMENT",
            file_path=result["file_path"],
        )

        task.task_output = {
            "format_version": format_version,
            "source": source,
            "file_path": result["file_path"],
            "artifact_id": artifact["id"],
        }
        task.status = "done"
        db.commit()
        session_label = "Project-level"
        if source == "session" and session_id:
            from app.models.session import Session as SessionModel
            session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
            if session and session.title:
                session_label = f'Session "{session.title}"'
            else:
                session_label = f"Session #{session_id}"
        log_action(
            db=db,
            user_id=user_id,
            project_id=project_id,
            action="generated",
            entity="srs_document",
            entity_id=artifact["id"],
            details={
                "label": f"SRS Document {artifact['version']}",
                "extra": f"{format_version} format ({session_label})"
            }
        )
        db.commit()

        # ── 5. Notify ──────────────────────────────────────────────────────
        _notify_srs_members(
            db=db, project_id=project_id, session_id=session_id,
            notification_type="srs_generated",
            title="SRS Document Ready",
            message=f"Your SRS document ({format_version}) has been generated. [project_id:{project_id}] [format_version:{format_version}]"
                    + (f" [session_id:{session_id}]" if session_id else ""),
        )

    except Exception as exc:
        logger.exception("SRS background task %s failed", task_id)
        try:
            task.status = "failed"
            task.error_message = _sanitize_srs_error_message(str(exc))
            db.commit()
        except Exception:
            db.rollback()
        try:
            _notify_srs_members(
                db=db, project_id=project_id, session_id=session_id,
                notification_type="srs_generation_failed",
                title="SRS Generation Failed",
                message="SRS document generation failed.",
            )
        except Exception:
            logger.exception("Failed to send SRS failure notification for task %s", task_id)
    finally:
        db.close()


def _notify_srs_members(db, project_id, session_id, notification_type, title, message):
    project = db.query(Project).filter(Project.id == project_id).first()
    project_name = project.name if project else None

    session_title = None
    if session_id:
        from app.models.session import Session as SessionModel
        session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
        session_title = session.title if session else None
    context = ""
    if session_title and project_name:
        context = f" for session '{session_title}' in {project_name}"
    elif session_title:
        context = f" for session '{session_title}'"
    elif project_name:
        context = f" in {project_name}"
    if context:
        stripped = message.rstrip()
        if stripped.endswith("."):
            message = stripped[:-1] + context + "."
        else:
            message = stripped + context

    if session_id:
        memberships = db.query(SessionMembership).filter(
            SessionMembership.session_id == session_id
        ).all()
        user_ids = [m.user_id for m in memberships]
    else:
        from app.models.project_membership import ProjectMembership
        members = db.query(ProjectMembership).filter(ProjectMembership.project_id == project_id).all()
        user_ids = [m.user_id for m in members]

    for uid in user_ids:
        notification_service.create_notification(
            db, user_id=uid,
            notification_type=notification_type,
            title=title, message=message,
            actor_name="System", actor_email=None,
            project_id=project_id, project_name=project_name,
        )
    db.commit()
 